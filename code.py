from docx import Document
import re

def remove_between_keywords(paragraphs, start_keyword, end_keyword):
    cleaned_paragraphs = []
    skip = False

    for paragraph in paragraphs:
        text = paragraph.text
        if start_keyword in text:
            skip = True
        if not skip:
            cleaned_paragraphs.append(paragraph)
        if end_keyword in text:
            skip = False

    return cleaned_paragraphs

# 打开Word文档
doc = Document('11.docx')

# 获取所有段落
paragraphs = doc.paragraphs

# 删除“解题思路”和“综上所述”之间的内容
cleaned_paragraphs = remove_between_keywords(paragraphs, "解题思路", "综上所述")

# 创建一个新的文档
new_doc = Document()

# 将清理后的段落添加到新文档中
for paragraph in cleaned_paragraphs:
    new_doc.add_paragraph(paragraph.text)

# 保存新文档
new_doc.save('cleaned_11.docx')

print("处理完成，新文档已保存为 cleaned_11.docx")
